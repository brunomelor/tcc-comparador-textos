# -*- coding: utf-8 -*-
"""
Gerador do TCC - Desenvolvimento de sistema de comparação de conteúdo de textos
Bruno Lorencatto - Engenharia de Software - UNINTER
"""

from __future__ import annotations

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import logging
import os
import re


# =============================================================================
# CONFIGURAÇÕES E HELPERS
# =============================================================================

def configurar_pagina(doc: Document) -> None:
    """Configura margens e tamanho da página conforme ABNT."""
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    # Numeração de página no canto superior direito
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    run.font.size = Pt(10)
    run.font.name = 'Arial'
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run._element.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._element.append(fldChar2)
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar3)


def estilo_padrao(doc: Document) -> None:
    """Configura estilos base do documento."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(1.25)


def add_titulo_artigo(doc: Document) -> None:
    """Adiciona título do artigo centralizado."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(
        'DESENVOLVIMENTO DE SISTEMA DE COMPARAÇÃO DE CONTEÚDO DE TEXTOS'
    )
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Arial'


def _add_nota_rodape(doc: Document, paragrafo, texto_nota: str) -> None:
    """Insere uma marca de nota de rodapé numérica em um parágrafo e adiciona o conteúdo da nota."""
    run = paragrafo.add_run()
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    fn_ref = OxmlElement('w:footnoteReference')
    # ID será atribuído via contagem; aqui usamos numeração sequencial nas notas
    fn_ref.set(qn('w:id'), str(_next_footnote_id(doc)))
    rPr = OxmlElement('w:rPr')
    vert = OxmlElement('w:vertAlign')
    vert.set(qn('w:val'), 'superscript')
    rPr.append(vert)
    run._element.append(rPr)
    run._element.append(fn_ref)
    _register_footnote(doc, fn_ref.get(qn('w:id')), texto_nota)


_FOOTNOTE_COUNTER = {'value': 0}
_FOOTNOTE_TEXTS: list[tuple[str, str]] = []


def _next_footnote_id(_doc: Document) -> int:
    _FOOTNOTE_COUNTER['value'] += 1
    return _FOOTNOTE_COUNTER['value']


def _register_footnote(_doc: Document, fid: str, texto: str) -> None:
    _FOOTNOTE_TEXTS.append((fid, texto))


def _materializar_notas_rodape(doc: Document) -> None:
    """Cria o part de footnotes e injeta as notas registradas. Chamar no final, antes de salvar."""
    if not _FOOTNOTE_TEXTS:
        return
    from docx.opc.constants import CONTENT_TYPE, RELATIONSHIP_TYPE
    from docx.opc.packuri import PackURI
    from docx.opc.part import Part
    from lxml import etree

    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    W = nsmap['w']

    # Constrói XML do part de footnotes
    root = etree.Element(f'{{{W}}}footnotes', nsmap=nsmap)
    # Notas obrigatórias separator (-1) e continuationSeparator (0)
    for sep_id, sep_type in [(-1, 'separator'), (0, 'continuationSeparator')]:
        fn = etree.SubElement(root, f'{{{W}}}footnote', {f'{{{W}}}id': str(sep_id), f'{{{W}}}type': sep_type})
        p = etree.SubElement(fn, f'{{{W}}}p')
        r = etree.SubElement(p, f'{{{W}}}r')
        sep_el = etree.SubElement(r, f'{{{W}}}separator' if sep_type == 'separator' else f'{{{W}}}continuationSeparator')

    for fid, texto in _FOOTNOTE_TEXTS:
        fn = etree.SubElement(root, f'{{{W}}}footnote', {f'{{{W}}}id': str(fid)})
        p = etree.SubElement(fn, f'{{{W}}}p')
        pPr = etree.SubElement(p, f'{{{W}}}pPr')
        pStyle = etree.SubElement(pPr, f'{{{W}}}pStyle', {f'{{{W}}}val': 'FootnoteText'})
        # Marca numérica
        r0 = etree.SubElement(p, f'{{{W}}}r')
        rPr0 = etree.SubElement(r0, f'{{{W}}}rPr')
        etree.SubElement(rPr0, f'{{{W}}}rStyle', {f'{{{W}}}val': 'FootnoteReference'})
        etree.SubElement(r0, f'{{{W}}}footnoteRef')
        # Espaço
        rsp = etree.SubElement(p, f'{{{W}}}r')
        tsp = etree.SubElement(rsp, f'{{{W}}}t', {'{http://www.w3.org/XML/1998/namespace}space': 'preserve'})
        tsp.text = ' '
        # Texto da nota
        r = etree.SubElement(p, f'{{{W}}}r')
        rPr = etree.SubElement(r, f'{{{W}}}rPr')
        rFonts = etree.SubElement(rPr, f'{{{W}}}rFonts', {f'{{{W}}}ascii': 'Arial', f'{{{W}}}hAnsi': 'Arial'})
        sz = etree.SubElement(rPr, f'{{{W}}}sz', {f'{{{W}}}val': '20'})
        t = etree.SubElement(r, f'{{{W}}}t')
        t.text = texto

    xml_bytes = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)

    partname = PackURI('/word/footnotes.xml')
    footnotes_part = Part(partname, CONTENT_TYPE.WML_FOOTNOTES, xml_bytes, doc.part.package)
    doc.part.relate_to(footnotes_part, RELATIONSHIP_TYPE.FOOTNOTES)


def add_autores(doc: Document) -> None:
    """Adiciona informações dos autores conforme template institucional UNINTER."""
    # Aluno
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run('LORENÇATTO, Bruno Meinertz')
    run.font.size = Pt(12)
    run.font.name = 'Arial'
    _add_nota_rodape(doc, p,
        'RU 3802493 — Bacharelando em Engenharia de Software, '
        'Centro Universitário Internacional – UNINTER.')

    # Orientador
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.first_line_indent = Cm(0)
    p2.paragraph_format.space_after = Pt(24)
    run2 = p2.add_run('FERRARI, Allan Christian Krainski')
    run2.font.size = Pt(12)
    run2.font.name = 'Arial'
    _add_nota_rodape(doc, p2,
        'Professor orientador — Escola Superior Politécnica, '
        'Centro Universitário Internacional – UNINTER.')


def add_secao_primaria(doc: Document, texto: str, centralizado: bool = False) -> None:
    """Seção nível 1: MAIÚSCULA + NEGRITO (ex: 1 INTRODUÇÃO). Use centralizado=True para RESUMO/REFERÊNCIAS/APÊNDICE."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.first_line_indent = Cm(0)
    if centralizado:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Arial'


def add_secao_secundaria(doc: Document, texto: str) -> None:
    """Seção nível 2: MAIÚSCULA SEM NEGRITO (ex: 2.1 SEÇÃO)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(texto.upper())
    run.bold = False
    run.font.size = Pt(12)
    run.font.name = 'Arial'


def add_secao_terciaria(doc: Document, texto: str) -> None:
    """Seção nível 3: minúscula + negrito (ex: 2.1.1 Seção terciária)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Arial'


def add_paragrafo(doc: Document, texto: str, recuo: bool = True) -> None:
    """Adiciona parágrafo de texto normal. Aceita marcadores *texto* para itálico (termos estrangeiros)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if not recuo:
        p.paragraph_format.first_line_indent = Cm(0)
    parts = re.split(r'(\*[^*]+\*)', texto)
    for part in parts:
        if not part:
            continue
        is_italic = len(part) >= 3 and part.startswith('*') and part.endswith('*')
        run = p.add_run(part[1:-1] if is_italic else part)
        run.font.size = Pt(12)
        run.font.name = 'Arial'
        if is_italic:
            run.italic = True


def add_placeholder(doc: Document, texto: str) -> None:
    """Adiciona texto placeholder em vermelho."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(f'[{texto}]')
    run.font.size = Pt(12)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.italic = True


def add_resumo(doc: Document, texto: str, palavras_chave: str) -> None:
    """Adiciona seção de Resumo com formatação específica (Arial 10, simples, título centralizado)."""
    add_secao_primaria(doc, 'Resumo', centralizado=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.first_line_indent = Cm(0)
    for part in re.split(r'(\*[^*]+\*)', texto):
        if not part:
            continue
        is_italic = len(part) >= 3 and part.startswith('*') and part.endswith('*')
        run = p.add_run(part[1:-1] if is_italic else part)
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        if is_italic:
            run.italic = True

    p2 = doc.add_paragraph()
    p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p2.paragraph_format.first_line_indent = Cm(0)
    p2.paragraph_format.space_before = Pt(12)
    run_label = p2.add_run('Palavras-chave: ')
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_label.font.name = 'Arial'
    run_kw = p2.add_run(palavras_chave)
    run_kw.bold = False
    run_kw.font.size = Pt(10)
    run_kw.font.name = 'Arial'


def add_referencias_formatadas(doc: Document, referencias: list[str]) -> None:
    """Adiciona referências com formatação ABNT (espaçamento simples, títulos em negrito)."""
    add_secao_primaria(doc, 'Referências', centralizado=True)
    for ref in referencias:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(12)
        # Parse **bold** markup para títulos em negrito (ABNT)
        parts = re.split(r'(\*\*.*?\*\*)', ref)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                run = p.add_run(part)
            run.font.size = Pt(12)
            run.font.name = 'Arial'


def add_tabela(doc: Document, titulo: str, cabecalhos: list[str], linhas: list[list[str]], fonte: str = "Elaborado pelo autor (2026)") -> None:
    """Adiciona tabela com formatação ABNT (legenda acima, fonte abaixo)."""
    # Legenda acima
    p_titulo = doc.add_paragraph()
    p_titulo.paragraph_format.space_before = Pt(12)
    p_titulo.paragraph_format.first_line_indent = Cm(0)
    run_t = p_titulo.add_run(titulo)
    run_t.font.size = Pt(10)
    run_t.font.name = 'Arial'

    # Tabela
    table = doc.add_table(rows=1 + len(linhas), cols=len(cabecalhos))
    table.style = 'Table Grid'

    # Cabeçalho
    for i, cab in enumerate(cabecalhos):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cab)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Arial'

    # Linhas de dados
    for r, linha in enumerate(linhas):
        for c, valor in enumerate(linha):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for part in re.split(r'(\*[^*]+\*)', str(valor)):
                if not part:
                    continue
                is_italic = len(part) >= 3 and part.startswith('*') and part.endswith('*')
                run = p.add_run(part[1:-1] if is_italic else part)
                run.font.size = Pt(10)
                run.font.name = 'Arial'
                if is_italic:
                    run.italic = True

    # Fonte abaixo
    p_fonte = doc.add_paragraph()
    p_fonte.paragraph_format.first_line_indent = Cm(0)
    p_fonte.paragraph_format.space_after = Pt(12)
    run_f = p_fonte.add_run(f'Fonte: {fonte}')
    run_f.font.size = Pt(10)
    run_f.font.name = 'Arial'


def add_figura(doc: Document, titulo: str, caminho_img: str, largura_cm: float = 14,
               fonte: str = "Elaborado pelo autor (2026)") -> None:
    """Adiciona figura com legenda ABNT (descrição acima, fonte abaixo).
    Se o arquivo de imagem existir em caminho_img, insere a imagem;
    caso contrário, insere uma moldura placeholder com instrução."""
    # Legenda acima
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.paragraph_format.first_line_indent = Cm(0)
    p_titulo.paragraph_format.space_before = Pt(12)
    p_titulo.paragraph_format.space_after = Pt(6)
    run_t = p_titulo.add_run(titulo)
    run_t.font.size = Pt(10)
    run_t.font.name = 'Arial'

    # Imagem ou placeholder
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.first_line_indent = Cm(0)
    p_img.paragraph_format.space_after = Pt(6)
    if caminho_img and os.path.isfile(caminho_img):
        run_img = p_img.add_run()
        run_img.add_picture(caminho_img, width=Cm(largura_cm))
    else:
        run_ph = p_img.add_run(f'[ Inserir imagem: {caminho_img} ]')
        run_ph.font.size = Pt(10)
        run_ph.font.name = 'Arial'
        run_ph.italic = True
        run_ph.font.color.rgb = RGBColor(160, 160, 160)

    # Fonte abaixo
    p_fonte = doc.add_paragraph()
    p_fonte.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fonte.paragraph_format.first_line_indent = Cm(0)
    p_fonte.paragraph_format.space_after = Pt(12)
    run_f = p_fonte.add_run(f'Fonte: {fonte}')
    run_f.font.size = Pt(10)
    run_f.font.name = 'Arial'


# =============================================================================
# REFERÊNCIAS BIBLIOGRÁFICAS (ABNT)
# =============================================================================

REFERENCIAS = [
    'BIRD, S.; KLEIN, E.; LOPER, E. **Natural Language Processing with Python**. Sebastopol: O\'Reilly Media, 2009.',
    'GIL, A. C. **Como elaborar projetos de pesquisa**. 6. ed. São Paulo: Atlas, 2017.',
    'GOMAA, W. H.; FAHMY, A. A. A survey of text similarity approaches. **International Journal of Computer Applications**, v. 68, n. 13, p. 13-18, 2013.',
    'MANNING, C. D.; RAGHAVAN, P.; SCHÜTZE, H. **Introduction to Information Retrieval**. Cambridge: Cambridge University Press, 2008.',
    'MORAIS, M. de F.; BOIKO, T. J. P. Metodologia de Pesquisa: uma proposta de estrutura para pesquisas técnico-científicas em Engenharia de Produção. **VIII Encontro de Engenharia de Produção Agroindustrial**, v. 1, p. 1-12, 2013.',
    'ORENGO, V. M.; HUYCK, C. R. A stemming algorithm for the Portuguese language. In: INTERNATIONAL SYMPOSIUM ON STRING PROCESSING AND INFORMATION RETRIEVAL (SPIRE), 8., 2001, Laguna de San Rafael. **Proceedings** [...]. Laguna de San Rafael: [s. n.], 2001. p. 186-193.',
    'PRESSMAN, R. S.; MAXIM, B. R. **Engenharia de Software**: uma abordagem profissional. 9. ed. Porto Alegre: AMGH, 2021.',
    'PRODANOV, C. C.; FREITAS, E. C. **Metodologia do trabalho científico**: métodos e técnicas da pesquisa e do trabalho acadêmico. 2. ed. Novo Hamburgo: Feevale, 2013.',
    'RAMOS, J. Using TF-IDF to determine word relevance in document queries. In: FIRST INSTRUCTIONAL CONFERENCE ON MACHINE LEARNING, 2003, Piscataway. **Proceedings** [...]. Piscataway: [s. n.], 2003. p. 133-142.',
    'REAL, L.; FONSECA, E.; GONÇALO OLIVEIRA, H. The ASSIN 2 shared task: a quick overview. In: INTERNATIONAL CONFERENCE ON COMPUTATIONAL PROCESSING OF THE PORTUGUESE LANGUAGE (PROPOR), 14., 2020, Évora. **Proceedings** [...]. Cham: Springer, 2020. p. 406-412.',
    'REIMERS, N.; GUREVYCH, I. Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks. In: CONFERENCE ON EMPIRICAL METHODS IN NATURAL LANGUAGE PROCESSING (EMNLP), 2019, Hong Kong. **Proceedings** [...]. Hong Kong: Association for Computational Linguistics, 2019. p. 3982-3992.',
    'SALTON, G.; BUCKLEY, C. Term-weighting approaches in automatic text retrieval. **Information Processing & Management**, v. 24, n. 5, p. 513-523, 1988.',
    'SANTOS, D. Processamento de linguagem natural: uma apresentação. In: CONGRESSO DA SOCIEDADE BRASILEIRA DE COMPUTAÇÃO, 22., 2002, Florianópolis. **Anais** [...]. Florianópolis: SBC, 2002. p. 59-82.',
    'SOMMERVILLE, I. **Engenharia de Software**. 10. ed. São Paulo: Pearson, 2019.',
    'VIEIRA, R.; LIMA, V. L. S. Linguística computacional: princípios e aplicações. In: JORNADA DE ATUALIZAÇÃO EM INTELIGÊNCIA ARTIFICIAL, 20., 2001, São Leopoldo. **Anais** [...]. São Leopoldo: SBC, 2001. p. 47-88.',
    'VIJAYMEENA, M. K.; KAVITHA, K. A survey on similarity measures in text mining. **Machine Learning and Applications: An International Journal**, v. 3, n. 1, p. 19-28, 2016.',
    'WAZLAWICK, R. S. **Metodologia de Pesquisa para Ciência da Computação**. 2. ed. Rio de Janeiro: Elsevier, 2014.',
    'YIN, R. K. **Estudo de caso**: planejamento e métodos. 5. ed. Porto Alegre: Bookman, 2015.',
]


# =============================================================================
# SEÇÕES DO ARTIGO
# =============================================================================

def secao_resumo(doc: Document) -> None:
    """RESUMO."""
    add_resumo(
        doc,
        'A produção crescente de conteúdo digital em ambientes '
        'acadêmicos, jurídicos e corporativos aumenta a demanda por '
        'ferramentas que identifiquem relações entre documentos '
        'textuais. Este trabalho descreve o desenvolvimento de um '
        'sistema web de comparação de textos em português. O sistema '
        'combina TF-IDF, similaridade por cosseno e coeficiente de '
        'Jaccard a um mecanismo que preserva a forma original das '
        'palavras após a stemização, melhorando a leitura dos termos '
        'exibidos ao usuário. A relevância do trabalho está em reunir, '
        'em um único sistema funcional, transparência, validação '
        'empírica e usabilidade. A metodologia adota pesquisa aplicada, '
        'quantitativa e exploratória, conduzida como estudo de caso, '
        'com avaliação sobre 100 pares de sentenças do corpus ASSIN2. '
        'O sistema utiliza Python com FastAPI no *back-end*, NLTK e '
        'scikit-learn no pipeline de PLN, e React com TypeScript e '
        'Tailwind CSS no *front-end*. Os resultados mostram correlação '
        'de Pearson de 0,688 entre o cosseno calculado e o julgamento '
        'humano, com tempo mediano de 2 ms por par, indicando '
        'viabilidade para uso interativo e transparente.',
        'similaridade textual; processamento de linguagem natural; '
        'TF-IDF; interpretabilidade; ASSIN2.'
    )


def secao_introducao(doc: Document) -> None:
    """1 INTRODUÇÃO."""
    add_secao_primaria(doc, '1 Introdução')

    # Motivação / Problematização
    add_paragrafo(doc,
        'O crescimento exponencial da produção de conteúdo digital nas últimas '
        'décadas trouxe consigo o desafio de identificar relações e '
        'sobreposições entre documentos de forma automatizada. Em contextos '
        'acadêmicos, jurídicos e corporativos, a necessidade de verificar a '
        'originalidade de textos, detectar reutilização de conteúdo e medir '
        'a correlação entre documentos tornou-se uma demanda recorrente '
        '(GOMAA; FAHMY, 2013). Segundo Manning, Raghavan e Schütze (2008), a '
        'recuperação de informação e a análise de similaridade textual '
        'constituem pilares fundamentais da ciência da computação moderna, '
        'com aplicações que vão desde motores de busca até sistemas de '
        'detecção de plágio.'
    )

    add_paragrafo(doc,
        'Apesar da existência de ferramentas comerciais para comparação '
        'de textos, como Turnitin e Copyscape, e de bibliotecas de '
        'código aberto voltadas a desenvolvedores, as soluções '
        'proprietárias geralmente operam como caixa-preta, e as '
        'bibliotecas abertas exigem do usuário conhecimento técnico '
        'para serem utilizadas. Além disso, soluções voltadas ao '
        'português frequentemente carecem de tratamento adequado das '
        'particularidades morfológicas da língua, como a riqueza de '
        'flexões verbais e a variação de sufixos (SANTOS, 2002). Esse '
        'cenário evidencia a necessidade de sistemas que mostrem ao '
        'usuário não apenas o índice de similaridade, mas também os '
        'termos que sustentam esse índice, de forma direta e '
        'compreensível.'
    )

    # Pergunta de pesquisa
    add_paragrafo(doc,
        'Diante desse contexto, a pergunta que orienta este trabalho é: '
        'em que medida a combinação de TF-IDF, similaridade por cosseno '
        'e mapeamento entre radicais e palavras originais produz um '
        'índice de similaridade coerente com o julgamento humano em '
        'textos em português?'
    )

    # Objetivo geral
    add_paragrafo(doc,
        'O objetivo geral deste trabalho é desenvolver um sistema web '
        'de comparação de conteúdo de textos em português. O sistema '
        'integra TF-IDF, similaridade por cosseno e coeficiente de '
        'Jaccard a um mecanismo de mapeamento entre radicais e palavras '
        'originais, e seus resultados são comparados ao julgamento '
        'humano de um corpus público.'
    )

    # Objetivos específicos
    add_paragrafo(doc,
        'Para alcançar o objetivo geral, são definidos os seguintes '
        'objetivos específicos: (i) investigar as principais técnicas '
        'de similaridade textual disponíveis na literatura, identificando '
        'suas vantagens e limitações; (ii) projetar e implementar um '
        'sistema web em arquitetura cliente-servidor que integre '
        'pré-processamento adequado ao português, vetorização TF-IDF, '
        'cálculo das métricas de similaridade por cosseno e Jaccard, '
        'além de mapeamento reverso stem-palavra para apresentação '
        'interpretável dos termos compartilhados; (iii) validar '
        'empiricamente as métricas implementadas sobre o corpus ASSIN2 '
        '(REAL; FONSECA; GONÇALO OLIVEIRA, 2020), medindo a '
        'correlação entre os escores calculados e o julgamento humano '
        'de similaridade; e (iv) desenvolver e apresentar a interface '
        'web (*front-end*) que permite ao usuário carregar documentos '
        'e visualizar os resultados de forma transparente.'
    )

    # Justificativa
    add_paragrafo(doc,
        'A relevância do trabalho está em três aspectos. Primeiro, o '
        'sistema integra a um pipeline clássico de TF-IDF um mecanismo '
        'que recupera a forma original das palavras após a stemização, '
        'tornando os termos apresentados ao usuário mais legíveis do que '
        'os radicais produzidos pelas bibliotecas tradicionais (BIRD; '
        'KLEIN; LOPER, 2009). Segundo, o sistema é avaliado sobre um '
        'corpus público em português (ASSIN2), comparando seus resultados '
        'com o julgamento humano de similaridade. Terceiro, o trabalho '
        'documenta as decisões de engenharia tomadas em cada etapa do '
        'desenvolvimento, podendo servir de referência para projetos '
        'semelhantes.'
    )

    # Estrutura do trabalho
    add_paragrafo(doc,
        'Este artigo está organizado da seguinte forma: a Seção 2 '
        'traz a fundamentação teórica, abordando conceitos de PLN, '
        'representação vetorial de textos com as formalizações de '
        'TF-IDF, cosseno e Jaccard, métricas de similaridade e engenharia '
        'de software para sistemas web. A Seção 3 descreve a metodologia '
        'adotada, classificando a pesquisa, definindo o procedimento de '
        'avaliação empírica sobre o corpus ASSIN2 e as tecnologias '
        'selecionadas. A Seção 4 expõe os resultados e discussões, '
        'detalhando a arquitetura do sistema, a implementação do módulo '
        'de processamento, a interface desenvolvida e a validação '
        'empírica realizada. A Seção 5 traz as considerações finais e '
        'sugestões para trabalhos futuros. Por fim, o Apêndice A traz '
        'detalhes da avaliação realizada, permitindo que o experimento '
        'seja repetido.'
    )


def secao_fundamentacao(doc: Document) -> None:
    """2 FUNDAMENTAÇÃO TEÓRICA."""
    add_secao_primaria(doc, '2 Similaridade textual e engenharia de sistemas')

    add_paragrafo(doc,
        'A comparação automatizada de textos constitui um dos desafios centrais da '
        'ciência da computação contemporânea, situando-se na interseção entre o '
        'PLN, a recuperação de informação e a '
        'engenharia de software (MANNING; RAGHAVAN; SCHÜTZE, 2008). Esta seção '
        'reúne os conceitos teóricos que sustentam o '
        'desenvolvimento de um sistema de comparação de conteúdo textual, '
        'abrangendo desde as técnicas clássicas de representação de documentos '
        'até as abordagens modernas baseadas em aprendizado profundo.'
    )

    # --- 2.1 ---
    add_secao_secundaria(doc, '2.1 Processamento de Linguagem Natural')

    add_paragrafo(doc,
        'Como ramo da inteligência artificial, o PLN reúne métodos que '
        'permitem aos computadores trabalhar com a linguagem humana em '
        'tarefas como leitura, interpretação e geração de texto. Vieira '
        'e Lima (2001) descrevem o campo como um conjunto de técnicas '
        'com base teórica que tornam possível representar e analisar a '
        'linguagem natural por meios computacionais, abrindo caminho '
        'para que sistemas processem dados textuais de modo útil.'
    )

    add_paragrafo(doc,
        'O tratamento de um texto em linguagem natural costuma envolver '
        'quatro etapas elementares. A primeira é a tokenização, em que o '
        'texto é dividido em palavras ou sentenças. A segunda consiste '
        'na remoção de *stopwords*, ou seja, de palavras frequentes mas '
        'com pouco valor semântico, como artigos e preposições. A '
        'terceira é a stemização, que reduz cada palavra a seu radical; '
        'no caso do português, destaca-se o algoritmo RSLP de Orengo e '
        'Huyck (2001), criado para lidar com as particularidades '
        'morfológicas do idioma. Por fim, a lematização busca a forma '
        'canônica das palavras e tende a preservar melhor o significado '
        'original (BIRD; KLEIN; LOPER, 2009).'
    )

    add_paragrafo(doc,
        'No caso específico do português, Santos (2002) chama atenção '
        'para desafios próprios do idioma, como a grande variedade de '
        'flexões verbais e a ambiguidade léxica, que exigem soluções '
        'distintas daquelas adotadas para o inglês. Esses fatores '
        'impactam diretamente a qualidade obtida por sistemas de '
        'comparação textual aplicados a documentos em português.'
    )

    # --- 2.2 ---
    add_secao_secundaria(doc, '2.2 Representação vetorial de textos')

    add_paragrafo(doc,
        'Para que um texto possa ser comparado por algoritmos, é preciso '
        'transformá-lo em uma representação numérica. O modelo de espaço '
        'vetorial, originalmente formalizado por Salton e Buckley (1988), '
        'tornou-se a base de boa parte das abordagens posteriores ao '
        'converter cada documento em um vetor situado em um espaço '
        'multidimensional.'
    )

    add_paragrafo(doc,
        'Uma das representações mais simples é o modelo *Bag of Words* '
        '(BoW), em que cada documento é descrito por um vetor com a '
        'frequência de seus termos, sem considerar a ordem das palavras '
        'ou a estrutura gramatical do texto. Apesar da facilidade de '
        'implementação, o BoW deixa de capturar o contexto semântico dos '
        'termos (MANNING; RAGHAVAN; SCHÜTZE, 2008).'
    )

    add_paragrafo(doc,
        'A técnica TF-IDF (*Term Frequency – Inverse Document Frequency*) '
        'estende o BoW ao atribuir pesos diferentes aos termos. O '
        'componente TF reflete o quanto um termo aparece em um documento; '
        'o IDF, por sua vez, reduz o peso de termos comuns ao conjunto, '
        'destacando aqueles que mais ajudam a distinguir um documento dos '
        'demais (RAMOS, 2003). De maneira formal, o peso TF-IDF de um '
        'termo *t* em um documento *d* dentro de uma coleção de *N* '
        'documentos é dado pela Equação 1, na qual tf(*t*,*d*) representa '
        'a frequência bruta do termo no documento e df(*t*) corresponde '
        'à quantidade de documentos da coleção que contêm aquele termo.'
    )

    add_figura(doc,
        'Equação 1 – Fórmula do peso TF-IDF',
        os.path.join(os.path.dirname(os.path.abspath(__file__)), 'docs', 'img', 'eq-tfidf.png'),
        largura_cm=10,
        fonte='Adaptado de Ramos (2003) e Manning, Raghavan e Schütze (2008)',
    )

    add_paragrafo(doc,
        'Mesmo com o surgimento de abordagens mais recentes, o TF-IDF '
        'segue presente em boa parte dos sistemas de recuperação de '
        'informação e de comparação textual, com bons resultados quando '
        'combinado a métricas como a similaridade por cosseno (VIJAYMEENA; '
        'KAVITHA, 2016).'
    )

    add_paragrafo(doc,
        'Já as abordagens mais modernas substituem a contagem de termos '
        'pelos chamados *embeddings* — vetores densos que codificam '
        'relações de significado entre palavras. Um exemplo é o '
        '*Sentence-BERT* (SBERT), proposto por Reimers e Gurevych '
        '(2019): trata-se de uma variante baseada em redes siamesas '
        'que gera *embeddings* no nível da sentença e diminui '
        'substancialmente o tempo de comparação entre pares de textos.'
    )

    # --- 2.3 ---
    add_secao_secundaria(doc, '2.3 Métricas de similaridade textual')

    add_paragrafo(doc,
        'As métricas de similaridade textual servem para quantificar o '
        'quanto dois documentos, ou dois trechos de texto, se parecem '
        'entre si. Em sua classificação, Gomaa e Fahmy (2013) organizam '
        'as métricas em três grandes grupos: as que comparam cadeias de '
        'caracteres (*string-based*), as que dependem de um corpus de '
        'apoio (*corpus-based*) e as que recorrem a bases de conhecimento '
        '(*knowledge-based*).'
    )

    add_paragrafo(doc,
        'Entre essas métricas, a similaridade por cosseno é a mais '
        'comum em sistemas de comparação textual. Ela é obtida pelo '
        'cosseno do ângulo formado entre os vetores que representam os '
        'documentos no espaço multidimensional, resultando em um valor '
        'entre 0 (ausência de similaridade) e 1 (textos idênticos). Sua '
        'principal vantagem é desconsiderar o tamanho do documento, uma '
        'vez que leva em conta apenas a direção dos vetores, e não suas '
        'magnitudes (MANNING; RAGHAVAN; SCHÜTZE, 2008). Para dois '
        'vetores *A* e *B* de dimensão *n*, a expressão correspondente '
        'aparece na Equação 2.'
    )

    add_figura(doc,
        'Equação 2 – Similaridade por cosseno entre dois vetores',
        os.path.join(os.path.dirname(os.path.abspath(__file__)), 'docs', 'img', 'eq-cosseno.png'),
        largura_cm=12,
        fonte='Adaptado de Manning, Raghavan e Schütze (2008)',
    )

    add_paragrafo(doc,
        'Outra métrica frequentemente empregada é o coeficiente de '
        'Jaccard, que avalia o quanto dois conjuntos *A* e *B* têm em '
        'comum por meio da razão entre o tamanho da interseção e o '
        'tamanho da união (Equação 3). Apesar de fácil de implementar e '
        'de interpretar, o Jaccard não captura aspectos semânticos: ele '
        'compara apenas os termos que de fato aparecem nos textos '
        '(VIJAYMEENA; KAVITHA, 2016).'
    )

    add_figura(doc,
        'Equação 3 – Coeficiente de Jaccard entre dois conjuntos',
        os.path.join(os.path.dirname(os.path.abspath(__file__)), 'docs', 'img', 'eq-jaccard.png'),
        largura_cm=7,
        fonte='Adaptado de Vijaymeena e Kavitha (2016)',
    )

    add_paragrafo(doc,
        'Quando aplicadas sobre vetores gerados por modelos como BERT ou '
        'SBERT, métricas baseadas em *embeddings* conseguem identificar '
        'similaridade entre textos mesmo quando estes empregam '
        'vocabulários diferentes. Reimers e Gurevych (2019) relatam '
        'ganhos expressivos de desempenho ao introduzir o SBERT: o '
        'tempo necessário para comparar 10.000 sentenças cai de cerca '
        'de 65 horas, no BERT original, para aproximadamente 5 segundos.'
    )

    # --- 2.4 ---
    add_secao_secundaria(doc, '2.4 Sistemas de detecção de similaridade')

    add_paragrafo(doc,
        'Os sistemas de detecção de similaridade textual encontram '
        'aplicação em diferentes contextos, sendo a verificação de '
        'plágio o uso mais conhecido. Em linhas gerais, esses sistemas '
        'seguem um fluxo composto pelo pré-processamento do texto, pela '
        'sua divisão em unidades passíveis de comparação, pela '
        'representação vetorial dessas unidades e, por fim, pela '
        'aplicação de métricas que comparam o texto a uma base de '
        'referência (GOMAA; FAHMY, 2013).'
    )

    add_paragrafo(doc,
        'No mercado, ferramentas como Turnitin e Copyscape recorrem a '
        'bases de dados amplas para realizar as comparações. No meio '
        'acadêmico, sistemas como o MOSS (*Measure of Software '
        'Similarity*), desenvolvido na Universidade de Stanford, '
        'destacam-se na comparação de código-fonte. Para documentos '
        'textuais, a estratégia mais frequente combina TF-IDF e '
        'similaridade por cosseno, podendo ser reforçada por técnicas '
        'de *fingerprinting* baseadas em n-gramas (MANNING; RAGHAVAN; '
        'SCHÜTZE, 2008).'
    )

    add_paragrafo(doc,
        'A escolha da técnica, segundo Vijaymeena e Kavitha (2016), '
        'depende do problema a ser resolvido: para a detecção de '
        'plágio literal, métricas baseadas em cadeias de caracteres '
        'tendem a funcionar melhor; já a identificação de paráfrases ou '
        'reescritas de conteúdo costuma exigir métricas com viés '
        'semântico. Em diversos casos, a combinação de mais de uma '
        'técnica produz resultados mais consistentes.'
    )

    # --- 2.5 ---
    add_secao_secundaria(doc, '2.5 Engenharia de software para sistemas web')

    add_paragrafo(doc,
        'A construção de um sistema de comparação textual no formato de '
        'aplicação web envolve a aplicação de princípios consolidados da '
        'engenharia de software. Pressman e Maxim (2021) definem a área '
        'como o conjunto de métodos, ferramentas e procedimentos que '
        'permitem controlar o processo de desenvolvimento e produzir '
        'software de qualidade de forma sustentável.'
    )

    add_paragrafo(doc,
        'Aplicações web, em particular, apresentam características que '
        'influenciam diretamente seu desenvolvimento: interfaces de '
        'usuário intuitivas, exigências de desempenho em tempo de '
        'execução e a separação entre as camadas de apresentação, lógica '
        'de negócios e dados (SOMMERVILLE, 2019). O padrão '
        'cliente-servidor, predominante nesse contexto, possibilita que '
        'tarefas computacionalmente caras, como o cálculo de similaridade '
        'textual, sejam realizadas no servidor, enquanto o *front-end* '
        'concentra-se na interação com o usuário.'
    )

    add_paragrafo(doc,
        'Quando combinadas a tarefas de PLN, as aplicações web enfrentam '
        'desafios próprios. O pré-processamento e a vetorização tendem '
        'a ser operações intensivas em CPU, e bibliotecas como NLTK e '
        'scikit-learn são síncronas por natureza, o que limita o ganho '
        'obtido por servidores assíncronos quando o trabalho é '
        'computacionalmente pesado (BIRD; KLEIN; LOPER, 2009). É '
        'preciso, portanto, dimensionar limites de entrada, prever '
        'tempos de resposta perceptíveis ao usuário e, quando o volume '
        'de requisições crescer, considerar estratégias como filas de '
        'processamento ou *cache* de resultados.'
    )

    add_paragrafo(doc,
        'Na definição da plataforma tecnológica é preciso ponderar '
        'aspectos como a oferta de bibliotecas de PLN, o desempenho '
        'esperado no processamento de textos e a facilidade de '
        'integração entre *back-end* e *front-end*. Nesse cenário, a '
        'linguagem Python ocupa posição de destaque graças ao seu '
        'ecossistema voltado ao processamento de linguagem natural — '
        'com bibliotecas como NLTK e spaCy — e a *frameworks* '
        'científicos como o scikit-learn, que já implementam algoritmos '
        'de TF-IDF e métricas de similaridade prontos para uso '
        '(BIRD; KLEIN; LOPER, 2009).'
    )


def secao_metodologia(doc: Document) -> None:
    """3 METODOLOGIA."""
    add_secao_primaria(doc, '3 Metodologia')

    add_paragrafo(doc,
        'Esta seção descreve os procedimentos metodológicos adotados para o '
        'desenvolvimento do sistema de comparação de conteúdo de textos, '
        'abrangendo a classificação da pesquisa, as etapas de desenvolvimento '
        'e as ferramentas tecnológicas selecionadas.'
    )

    # --- 3.1 ---
    add_secao_secundaria(doc, '3.1 Classificação da pesquisa')

    add_paragrafo(doc,
        'A classificação de uma pesquisa, segundo Morais e Boiko (2013), '
        'organiza-se em quatro grupos: natureza, abordagem do problema, '
        'objetivos e procedimentos técnicos. Esse modelo, também '
        'sistematizado por Prodanov e Freitas (2013) e Wazlawick (2014), '
        'serve de base para a classificação a seguir.'
    )

    add_paragrafo(doc,
        'Quanto à natureza, este trabalho enquadra-se como pesquisa '
        'aplicada, uma vez que o conhecimento produzido tem destino '
        'prático imediato: oferecer um meio automatizado para comparar '
        'o conteúdo de dois documentos textuais (GIL, 2017; PRODANOV; '
        'FREITAS, 2013). O resultado final é um sistema funcional que '
        'calcula e exibe o índice de correlação entre os textos '
        'analisados.'
    )

    add_paragrafo(doc,
        'Quanto à abordagem, o trabalho adota o enfoque quantitativo, '
        'pois os resultados da comparação textual são expressos por '
        'métricas numéricas — como os índices de similaridade por '
        'cosseno e Jaccard — que permitem medir de forma objetiva o '
        'grau de correlação entre os documentos (GIL, 2017).'
    )

    add_paragrafo(doc,
        'Em relação aos objetivos, a pesquisa caracteriza-se como '
        'exploratória, uma vez que investiga diferentes técnicas de '
        'similaridade textual e suas implementações, buscando '
        'identificar a abordagem mais adequada ao sistema proposto. '
        'Quanto aos procedimentos, combinam-se a pesquisa bibliográfica '
        '— para embasar as decisões técnicas — e o estudo de caso '
        '(YIN, 2015; WAZLAWICK, 2014), aqui representado pela própria '
        'construção do sistema. Para avaliar os resultados, as métricas '
        'implementadas são aplicadas a 100 pares de sentenças do corpus '
        'ASSIN2 (REAL; FONSECA; GONÇALO OLIVEIRA, 2020), e os escores '
        'calculados são comparados ao julgamento humano por meio das '
        'correlações de Pearson e Spearman, do erro absoluto médio (MAE) '
        'e do erro quadrático médio (MSE).'
    )

    # --- 3.2 ---
    add_secao_secundaria(doc, '3.2 Etapas do desenvolvimento')

    add_paragrafo(doc,
        'O desenvolvimento do sistema segue uma abordagem incremental, '
        'alinhada ao modelo de processo evolutivo descrito por Pressman e '
        'Maxim (2021), organizada nas seguintes etapas:'
    )

    add_paragrafo(doc,
        '1) Levantamento bibliográfico: revisão da literatura sobre técnicas '
        'de PLN, métricas de similaridade textual '
        'e arquiteturas de sistemas web, com o objetivo de fundamentar as '
        'decisões de projeto. As buscas são realizadas no Google Acadêmico, '
        'IEEE Xplore e SciELO, utilizando as palavras-chave: "similaridade '
        'textual", "text similarity", "TF-IDF", "cosine similarity", '
        '"comparação de documentos" e "NLP".'
    )

    add_paragrafo(doc,
        '2) Definição dos requisitos: especificação das funcionalidades do '
        'sistema, incluindo: *upload* de dois arquivos de texto, '
        'pré-processamento automático (tokenização, remoção de *stopwords*, '
        'stemização), cálculo de múltiplas métricas de similaridade e '
        'apresentação dos resultados em interface gráfica.'
    )

    add_paragrafo(doc,
        '3) Escolha da plataforma tecnológica: seleção da linguagem de '
        'programação, *frameworks* de *back-end* e *front-end*, e bibliotecas de '
        'PLN, considerando critérios como maturidade do ecossistema, '
        'disponibilidade de bibliotecas especializadas e facilidade de '
        'implantação.'
    )

    add_paragrafo(doc,
        '4) Projeto da arquitetura: definição da arquitetura do sistema, '
        'incluindo a separação entre camadas de apresentação (*front-end*), '
        'lógica de negócios (*back-end*) e processamento de PLN, seguindo a '
        'arquitetura cliente-servidor.'
    )

    add_paragrafo(doc,
        '5) Implementação do módulo de processamento: desenvolvimento do '
        'núcleo do sistema responsável pelo pré-processamento dos textos, '
        'geração de representações vetoriais (TF-IDF) e cálculo das métricas '
        'de similaridade (cosseno e Jaccard).'
    )

    add_paragrafo(doc,
        '6) Desenvolvimento do *front-end*: construção da interface web que '
        'permite ao usuário carregar documentos, visualizar os índices de '
        'correlação e explorar os trechos de maior e menor similaridade entre '
        'os textos comparados.'
    )

    add_paragrafo(doc,
        '7) Testes e validação: execução de duas frentes de avaliação. '
        '(a) Testes automatizados unitários do módulo de processamento '
        'utilizando *pytest*, cobrindo cenários extremos (textos idênticos, '
        'textos sem sobreposição e textos parcialmente similares) e '
        'garantindo regressão. (b) Validação empírica das métricas sobre '
        '100 pares amostrados aleatoriamente do *split* de validação do '
        'corpus ASSIN2 (REAL; FONSECA; GONÇALO OLIVEIRA, 2020), corpus '
        'anotado por avaliadores humanos em escala de 1 a 5 para '
        'similaridade semântica. Os escores humanos são normalizados '
        'linearmente para o intervalo [0,1] antes da comparação com a '
        'saída do sistema, e a comparação é feita pelas correlações de '
        'Pearson e Spearman e pelos erros absoluto e quadrático médios. '
        'A semente aleatória é fixada (42) para permitir que o '
        'experimento seja repetido.'
    )

    # --- 3.3 ---
    add_secao_secundaria(doc, '3.3 Ferramentas e tecnologias')

    add_paragrafo(doc,
        'A linguagem Python é adotada como tecnologia principal do *back-end*, '
        'pelo amplo ecossistema de bibliotecas voltadas ao PLN. '
        'As principais bibliotecas utilizadas são: '
        'NLTK (*Natural Language Toolkit*), para tokenização, remoção de '
        '*stopwords* e stemização; scikit-learn, para vetorização TF-IDF e '
        'cálculo de similaridade por cosseno; e FastAPI, como *framework* '
        'assíncrono para exposição da API REST, com geração automática '
        'de documentação interativa (BIRD; KLEIN; LOPER, 2009).'
    )

    add_paragrafo(doc,
        'Para o *front-end*, utiliza-se o *framework* React com TypeScript, '
        'que proporciona tipagem estática e maior segurança no '
        'desenvolvimento de componentes reutilizáveis. A estilização é '
        'realizada com Tailwind CSS, um *framework* CSS utilitário que '
        'permite construir interfaces responsivas de forma ágil. A '
        'ferramenta de *build* Vite é empregada para o empacotamento e '
        'servimento da aplicação em ambiente de desenvolvimento. A '
        'comunicação entre *front-end* e *back-end* ocorre por meio de '
        'requisições HTTP à API REST, seguindo os princípios da '
        'arquitetura cliente-servidor (SOMMERVILLE, 2019).'
    )

    add_paragrafo(doc,
        'O controle de versão do código-fonte é realizado com Git, e o '
        'ambiente de desenvolvimento utiliza o editor Visual Studio Code. '
        'Para a documentação e gestão do projeto, são empregadas práticas '
        'da engenharia de software alinhadas ao modelo incremental, '
        'permitindo entregas parciais e refinamentos progressivos '
        '(PRESSMAN; MAXIM, 2021).'
    )


def secao_resultados(doc: Document) -> None:
    """4 RESULTADOS E DISCUSSÕES."""
    add_secao_primaria(doc, '4 Resultados e Discussões')

    add_paragrafo(doc,
        'Esta seção descreve os resultados obtidos em cada etapa '
        'na metodologia, seguindo a mesma ordem cronológica. São discutidos os '
        'resultados do levantamento bibliográfico, as decisões arquiteturais, os '
        'resultados da implementação do sistema e os testes de validação '
        'realizados.'
    )

    # --- 4.1 ---
    add_secao_secundaria(doc, '4.1 Resultados do levantamento bibliográfico')

    add_paragrafo(doc,
        'O levantamento bibliográfico realizado nas bases Google Acadêmico, '
        'IEEE Xplore e SciELO resultou na identificação de três abordagens '
        'principais para a comparação de conteúdo textual: abordagens baseadas '
        'em frequência de termos (TF-IDF), abordagens baseadas em similaridade '
        'de conjuntos (Jaccard) e abordagens baseadas em *embeddings* semânticos '
        '(BERT/SBERT). Conforme a taxonomia proposta por Gomaa e Fahmy (2013), '
        'essas técnicas pertencem, respectivamente, às categorias *corpus-based*, '
        '*string-based* e *knowledge-based*. A análise da literatura permitiu '
        'identificar as vantagens e limitações de cada técnica, conforme '
        'sintetizado no Quadro 1.'
    )

    add_tabela(doc,
        'Quadro 1 – Comparativo das técnicas de similaridade textual',
        ['Técnica', 'Tipo', 'Vantagens', 'Limitações'],
        [
            ['TF-IDF + Cosseno', 'Estatística',
             'Rápida, interpretável, eficaz para textos longos',
             'Não captura semântica'],
            ['Jaccard', 'Conjuntos',
             'Simples, intuitiva',
             'Ignora frequência e semântica'],
            ['SBERT', 'Embeddings',
             'Captura similaridade semântica',
             'Requer modelo pré-treinado, maior custo computacional'],
        ],
        'Elaborado pelo autor com base em Gomaa e Fahmy (2013) e Reimers e Gurevych (2019)'
    )

    add_paragrafo(doc,
        'A análise comparativa evidenciou que a combinação de TF-IDF com '
        'similaridade por cosseno oferece o melhor equilíbrio entre '
        'desempenho computacional e qualidade dos resultados para a '
        'comparação de documentos completos, como apontam Vijaymeena e '
        'Kavitha (2016) e Manning, Raghavan e Schütze (2008) ao tratar '
        'da eficácia do TF-IDF em tarefas de recuperação de informação. '
        'O Jaccard, por não considerar a frequência dos termos '
        '(VIJAYMEENA; KAVITHA, 2016), '
        'complementa a análise ao fornecer uma perspectiva baseada na '
        'sobreposição vocabular. Já as abordagens baseadas em *embeddings*, '
        'como o SBERT, apresentam resultados superiores na identificação '
        'de paráfrases (REIMERS; GUREVYCH, 2019), porém demandam recursos '
        'computacionais significativamente maiores.'
    )

    # --- 4.2 ---
    add_secao_secundaria(doc, '4.2 Definição da arquitetura e tecnologias')

    add_paragrafo(doc,
        'Com base nos requisitos levantados e na análise bibliográfica, '
        'define-se a arquitetura do sistema segundo o padrão cliente-servidor, '
        'modelo predominante no desenvolvimento web conforme Sommerville '
        '(2019), com separação clara entre as camadas de apresentação e '
        'processamento. A linguagem Python é selecionada para o *back-end* '
        'pelo ecossistema consolidado de bibliotecas de PLN '
        '(BIRD; KLEIN; LOPER, 2009) e da facilidade de prototipação. O '
        '*framework* FastAPI é adotado para a construção da API REST, por '
        'seu suporte nativo a operações assíncronas e geração automática '
        'de documentação interativa via Swagger UI, acessível pelo '
        '*endpoint* /docs. A API expõe dois *endpoints*: POST /api/compare, '
        'que recebe dois textos e retorna as métricas de similaridade, e '
        'GET /api/health, que permite verificar a disponibilidade do '
        'serviço. A comunicação entre *front-end* e *back-end* é habilitada '
        'por meio de CORS (*Cross-Origin Resource Sharing*), e os textos '
        'de entrada são validados com limite máximo de 500.000 caracteres '
        'para evitar sobrecarga do servidor.'
    )

    add_paragrafo(doc,
        'A Figura 1 sintetiza a arquitetura do sistema em três camadas: '
        'apresentação (cliente), API e processamento de PLN. A camada de '
        'apresentação concentra os componentes React responsáveis pela '
        'interação com o usuário; a camada de API expõe os *endpoints* '
        'REST do FastAPI; e a camada de processamento concentra as etapas '
        'do *pipeline* de PLN, incluindo o mapeamento reverso stem-palavra '
        'destacado como contribuição central deste trabalho.'
    )

    add_figura(doc,
        'Figura 1 – Arquitetura em camadas do sistema',
        os.path.join(os.path.dirname(os.path.abspath(__file__)), 'docs', 'img', 'arquitetura.png'),
        largura_cm=15,
    )

    add_paragrafo(doc,
        'A seleção das tecnologias adotadas em cada camada seguiu critérios '
        'objetivos: maturidade do ecossistema, presença de bibliotecas '
        'especializadas, qualidade da documentação oficial, suporte a '
        'tipagem estática (quando aplicável) e facilidade de implantação. '
        'O Quadro 2 sintetiza o mapeamento entre os componentes e as '
        'tecnologias selecionadas, com a respectiva justificativa.'
    )

    add_tabela(doc,
        'Quadro 2 – Tecnologias adotadas por componente do sistema',
        ['Componente', 'Tecnologia', 'Justificativa'],
        [
            ['*Back-end* / API', 'Python + FastAPI',
             'Ecossistema de PLN maduro; API com documentação automática'],
            ['Pré-processamento', 'NLTK',
             'Tokenização, *stopwords* e stemização RSLP para português'],
            ['Vetorização', 'scikit-learn (TfidfVectorizer)',
             'Implementação otimizada de TF-IDF'],
            ['Similaridade', 'scikit-learn (cosine_similarity)',
             'Cálculo eficiente de similaridade por cosseno'],
            ['*Front-end*', 'React + TypeScript + Tailwind CSS',
             'Componentes tipados, estilização ágil e responsiva'],
            ['*Build*', 'Vite',
             'Empacotamento rápido com *Hot Module Replacement*'],
            ['Controle de versão', 'Git',
             'Rastreabilidade e versionamento do código'],
        ],
        'Elaborado pelo autor (2026)'
    )

    add_paragrafo(doc,
        'A escolha entre *frameworks* concorrentes para o *back-end* foi '
        'guiada por análise comparativa explícita. O Quadro 3 apresenta '
        'a matriz de decisão entre Flask, FastAPI e Django, considerando '
        'critérios relevantes ao escopo do projeto. FastAPI foi escolhido '
        'por combinar tipagem nativa via Pydantic, geração automática de '
        'documentação OpenAPI e desempenho competitivo, com curva de '
        'aprendizado adequada ao prazo de desenvolvimento.'
    )

    add_tabela(doc,
        'Quadro 3 – Matriz de decisão dos frameworks web para Python',
        ['Critério', 'Flask', 'FastAPI', 'Django'],
        [
            ['Tipagem nativa', 'Não', 'Sim (Pydantic)', 'Parcial'],
            ['Documentação automática', 'Não', 'Sim (OpenAPI)', 'Parcial'],
            ['Suporte assíncrono', 'Limitado', 'Nativo', 'Parcial'],
            ['Curva de aprendizado', 'Baixa', 'Baixa-Média', 'Alta'],
            ['Adequação a APIs REST simples', 'Alta', 'Alta', 'Média'],
            ['Decisão final', '—', 'Selecionado', '—'],
        ],
        'Elaborado pelo autor (2026), com base na documentação oficial dos frameworks'
    )

    # --- 4.3 ---
    add_secao_secundaria(doc, '4.3 Implementação do módulo de processamento')

    add_paragrafo(doc,
        'O módulo de processamento constitui o núcleo do sistema e opera em '
        'três etapas sequenciais: pré-processamento, vetorização e cálculo '
        'de similaridade.'
    )

    add_paragrafo(doc,
        'Na etapa de pré-processamento, cada documento de entrada passa por '
        'tokenização (segmentação em palavras), conversão para minúsculas, '
        'remoção de *stopwords* da língua portuguesa utilizando a lista '
        'disponível no NLTK, e stemização por meio do algoritmo RSLP '
        '(Removedor de Sufixos da Língua Portuguesa), proposto por Orengo '
        'e Huyck (2001) especificamente para o tratamento das '
        'particularidades morfológicas do português. Essas operações '
        'reduzem a dimensionalidade do vocabulário e eliminam variações '
        'morfológicas que não contribuem para a análise de similaridade.'
    )

    add_paragrafo(doc,
        'Na etapa de vetorização, os textos pré-processados são convertidos '
        'em vetores numéricos utilizando o TfidfVectorizer da biblioteca '
        'scikit-learn, que implementa a técnica TF-IDF conforme descrita '
        'por Ramos (2003). O vetorizador é ajustado ao corpus formado '
        'pelos dois documentos, gerando uma matriz TF-IDF na qual cada '
        'linha representa um documento e cada coluna representa um termo '
        'ponderado pela sua relevância relativa.'
    )

    add_paragrafo(doc,
        'Na etapa de cálculo de similaridade, são aplicadas duas métricas '
        'sobre os vetores gerados: a similaridade por cosseno (Equação 2), '
        'amplamente utilizada em sistemas de comparação textual (MANNING; '
        'RAGHAVAN; SCHÜTZE, 2008), que mede a proximidade direcional entre '
        'os vetores e resulta em um valor entre 0 e 1; e o coeficiente de '
        'Jaccard (Equação 3), calculado sobre os conjuntos de stems únicos '
        'de cada documento após o pré-processamento (GOMAA; FAHMY, 2013). '
        'O cálculo sobre stems, e não sobre as formas originais, garante '
        'coerência com a representação vetorial usada pelo cosseno e '
        'reduz a influência de variações morfológicas. O índice de '
        'correlação final é apresentado como um percentual, facilitando '
        'a interpretação pelo usuário.'
    )

    add_secao_terciaria(doc, '4.3.1 Mapeamento reverso stem-palavra')

    add_paragrafo(doc,
        'Para melhorar a interpretabilidade dos resultados, o módulo de '
        'processamento foi acrescido de um mecanismo de mapeamento '
        'reverso entre radicais (stems) e palavras originais. As '
        'bibliotecas usuais, como NLTK e scikit-learn (BIRD; KLEIN; LOPER, '
        '2009), realizam a stemização sem preservar a forma original das '
        'palavras, o que dificulta a leitura dos termos quando estes são '
        'apresentados ao usuário. No sistema proposto, durante a '
        'stemização, cada radical gerado é associado à forma original '
        'mais completa encontrada nos textos de entrada e armazenado em '
        'um dicionário auxiliar (variável `stem_to_word`). Ao apresentar '
        'os termos compartilhados, a interface utiliza esse dicionário '
        'para exibir formas legíveis (por exemplo, "processamento") em '
        'vez de radicais truncados (por exemplo, "process"). O '
        'procedimento é simples, não acrescenta custo computacional '
        'significativo e melhora a clareza do retorno entregue ao '
        'usuário, atendendo ao princípio de transparência destacado '
        'por Pressman e Maxim (2021).'
    )

    add_paragrafo(doc,
        'A Figura 2 ilustra o fluxo completo de uma requisição ao *endpoint* '
        'POST /api/compare em formato de diagrama de sequência. O fluxo '
        'envolve cinco atores: a interface React, a camada FastAPI, o '
        'componente TextComparer (nlp.py) e as bibliotecas externas '
        'NLTK e scikit-learn. O retorno consolida métricas, termos '
        'compartilhados e estatísticas de processamento em um único '
        'objeto JSON.'
    )

    add_figura(doc,
        'Figura 2 – Diagrama de sequência do POST /api/compare',
        os.path.join(os.path.dirname(os.path.abspath(__file__)), 'docs', 'img', 'sequencia.png'),
        largura_cm=15,
    )

    add_paragrafo(doc,
        'O sistema também coleta estatísticas de processamento — contagem '
        'de palavras de cada texto e tempo de execução em milissegundos — '
        'que são apresentadas ao usuário junto aos resultados, conferindo '
        'transparência ao funcionamento do sistema.'
    )

    # --- 4.4 ---
    add_secao_secundaria(doc, '4.4 Desenvolvimento do front-end')

    add_paragrafo(doc,
        'A interface web é construída com React e TypeScript e estilizada '
        'com Tailwind CSS, organizada em quatro componentes reutilizáveis: '
        'TextInput (entrada com contador de palavras em tempo real e '
        '*upload* de arquivos .txt, .md, .csv e .log), Results (apresentação '
        'das métricas com barras de progresso coloridas por faixa de '
        'intensidade — verde acima de 70%, amarelo entre 40% e 70%, '
        'vermelho abaixo de 40% — e estatísticas de processamento), '
        'SharedTerms (tabela com os quinze termos compartilhados mais '
        'relevantes, ordenados pelo peso TF-IDF médio, acompanhada de '
        'listas de termos exclusivos como etiquetas coloridas) e Spinner '
        '(indicador de carregamento). A aplicação adota tipagem estática '
        'estrita e tratamento explícito de estados de carregamento e erro, '
        'em conformidade com as boas práticas descritas por Pressman e '
        'Maxim (2021) para sistemas web interativos.'
    )

    add_paragrafo(doc,
        'A Figura 3 apresenta a tela inicial do sistema, com o cabeçalho '
        'identificando o nome da aplicação e a descrição das métricas '
        'empregadas, e as duas áreas de entrada de texto dispostas lado '
        'a lado. A comparação pode ser acionada pelo botão "Comparar" ou '
        'pelo atalho de teclado Ctrl+Enter.'
    )

    add_figura(doc,
        'Figura 3 – Tela inicial do sistema com áreas de entrada de texto',
        os.path.join(os.path.dirname(os.path.abspath(__file__)), 'docs', 'img', 'tela-inicial.png'),
        largura_cm=14,
    )

    add_paragrafo(doc,
        'A Figura 4 ilustra a tela de resultados, exibida após o '
        'processamento dos textos. A parte superior traz as duas '
        'métricas de similaridade em formato percentual, com barras de '
        'progresso coloridas. Logo abaixo são apresentadas as estatísticas '
        'de processamento (contagem de palavras e tempo de execução). Na '
        'parte inferior, o componente de termos compartilhados exibe a '
        'tabela com os quinze termos mais relevantes comuns aos dois '
        'documentos e duas listas de termos exclusivos. Os termos '
        'apresentados são as formas legíveis recuperadas pelo mapeamento '
        'reverso descrito na Seção 4.3.1 — e não os radicais (stems) — '
        'o que materializa, na camada de apresentação, a contribuição '
        'central deste trabalho.'
    )

    add_figura(doc,
        'Figura 4 – Tela de resultados com métricas, estatísticas e termos compartilhados',
        os.path.join(os.path.dirname(os.path.abspath(__file__)), 'docs', 'img', 'tela-resultados.png'),
        largura_cm=14,
    )

    # --- 4.5 ---
    add_secao_secundaria(doc, '4.5 Testes e validação')

    add_paragrafo(doc,
        'A validação do sistema foi conduzida em duas frentes complementares: '
        '(a) testes unitários automatizados do módulo de processamento, '
        'cobrindo cenários extremos e regressão; e (b) avaliação empírica '
        'das métricas sobre o corpus ASSIN2 (REAL; FONSECA; GONÇALO '
        'OLIVEIRA, 2020), confrontando os escores calculados com o '
        'julgamento humano anotado por avaliadores especialistas.'
    )

    add_secao_terciaria(doc, '4.5.1 Testes unitários automatizados')

    add_paragrafo(doc,
        'O arquivo `test_nlp.py` reúne os testes unitários do módulo de '
        'processamento, executados com o *framework pytest*. A suíte cobre: '
        'pré-processamento (remoção de *stopwords*, stemização e mapeamento '
        'stem-palavra), cálculo das métricas de similaridade (cosseno e '
        'Jaccard) e cenários-limite (textos idênticos, textos disjuntos, '
        'entradas vazias e textos acima do limite de 500.000 caracteres). '
        'Os testes asseguram a manutenção do comportamento esperado a '
        'cada alteração no código, contribuindo para a sustentabilidade '
        'do sistema conforme Pressman e Maxim (2021).'
    )

    add_secao_terciaria(doc, '4.5.2 Validação empírica sobre o corpus ASSIN2')

    add_paragrafo(doc,
        'Para a avaliação empírica foram amostrados aleatoriamente 100 '
        'pares de sentenças do *split* de validação do ASSIN2, corpus '
        'anotado por avaliadores humanos em escala de 1 (nenhuma '
        'similaridade) a 5 (similaridade máxima). Os escores humanos '
        'foram linearmente normalizados para o intervalo [0,1] e '
        'confrontados com a saída do sistema. Foram calculados quatro '
        'indicadores estatísticos para cada métrica: correlação de '
        'Pearson, correlação de Spearman, erro absoluto médio (MAE) e '
        'erro quadrático médio (MSE). O Quadro 4 sintetiza os resultados.'
    )

    add_tabela(doc,
        'Quadro 4 – Correlação das métricas com o julgamento humano (corpus ASSIN2, n = 100)',
        ['Métrica', 'Pearson', 'Spearman', 'MAE', 'MSE'],
        [
            ['Cosseno (TF-IDF)', '0,688', '0,682', '0,236', '0,084'],
            ['Jaccard (stems)', '0,678', '0,680', '0,248', '0,091'],
        ],
        'Elaborado pelo autor (2026) com base em REAL, FONSECA e GONÇALO OLIVEIRA (2020)'
    )

    add_paragrafo(doc,
        'Os dois testes de correlação apresentaram p-valor inferior a '
        '0,001, indicando que os resultados observados não decorrem do '
        'acaso. A similaridade por cosseno apresentou resultado '
        'ligeiramente melhor que o coeficiente de Jaccard, tanto em '
        'Pearson (0,688 contra 0,678) quanto em MAE (0,236 contra 0,248). '
        'O resultado é compatível com a literatura, que aponta o cosseno '
        'sobre TF-IDF como técnica mais robusta por considerar a '
        'frequência ponderada dos termos (VIJAYMEENA; KAVITHA, 2016). '
        'Cabe ressaltar que sistemas baseados em *embeddings* '
        'semânticos, como o SBERT, costumam alcançar correlações de '
        'Pearson em torno de 0,80 no mesmo corpus (REIMERS; GUREVYCH, '
        '2019), patamar superior ao obtido aqui — diferença esperada, '
        'visto que o TF-IDF é uma métrica léxica e não captura '
        'paráfrases. O tempo mediano de processamento por par foi de '
        '2 ms, com pico de 45 ms, demonstrando viabilidade para uso '
        'interativo.'
    )

    add_paragrafo(doc,
        'A Figura 5 traz o gráfico de dispersão entre os escores '
        'calculados pelo sistema e os escores humanos do ASSIN2, em que '
        'cada ponto representa um dos 100 pares avaliados. A linha '
        'tracejada representa a diagonal *y = x* (concordância perfeita). '
        'A nuvem de pontos exibe tendência positiva clara em ambas as '
        'métricas, com maior dispersão em valores intermediários — '
        'comportamento esperado por se tratar de métricas léxicas, '
        'limitadas em capturar similaridade semântica de paráfrases '
        '(REIMERS; GUREVYCH, 2019).'
    )

    add_figura(doc,
        'Figura 5 – Dispersão entre similaridade calculada e julgamento humano (ASSIN2, n = 100)',
        os.path.join(os.path.dirname(os.path.abspath(__file__)), 'docs', 'img', 'dispersao-assin2.png'),
        largura_cm=15,
        fonte='Elaborado pelo autor (2026) a partir do corpus ASSIN2 (REAL; FONSECA; GONÇALO OLIVEIRA, 2020)',
    )

    add_paragrafo(doc,
        'A interface web mostrou-se funcional e responsiva durante todos '
        'os testes manuais. O mapeamento reverso de stems para palavras '
        'originais contribuiu para a interpretabilidade percebida, '
        'permitindo que o usuário identifique os termos legíveis que '
        'sustentam o índice de correlação — necessidade destacada por '
        'Bird, Klein e Loper (2009) como fundamental para a usabilidade '
        'de sistemas de PLN. As estatísticas exibidas (contagem de '
        'palavras e tempo de processamento) conferem transparência ao '
        'funcionamento do sistema, princípio essencial da engenharia de '
        'software segundo Pressman e Maxim (2021).'
    )


def secao_consideracoes(doc: Document) -> None:
    """5 CONSIDERAÇÕES FINAIS."""
    add_secao_primaria(doc, '5 Considerações Finais')

    add_paragrafo(doc,
        'Este trabalho teve como objetivo desenvolver um sistema web de '
        'comparação de textos em português que combinasse técnicas '
        'clássicas de similaridade a um mecanismo capaz de preservar a '
        'forma original das palavras após a stemização. A pergunta de '
        'pesquisa que orientou o estudo foi respondida de forma '
        'afirmativa: a combinação de TF-IDF com similaridade por '
        'cosseno, somada ao mapeamento reverso entre radicais e '
        'palavras, produz índices coerentes com o julgamento humano '
        'em corpus público de português e permite que o usuário '
        'visualize os termos relevantes em sua forma legível.'
    )

    add_paragrafo(doc,
        'O objetivo específico (i), referente à investigação das técnicas '
        'de similaridade textual, foi alcançado por meio do levantamento '
        'bibliográfico, que identificou e comparou três abordagens '
        'principais: TF-IDF com similaridade por cosseno, coeficiente de '
        'Jaccard e *embeddings* semânticos (SBERT). A análise indicou que '
        'a combinação TF-IDF com cosseno oferece bom equilíbrio entre '
        'desempenho e qualidade para a comparação de documentos em '
        'português, conclusão confirmada pela validação sobre o ASSIN2.'
    )

    add_paragrafo(doc,
        'O objetivo específico (ii), relativo à concepção e à '
        'implementação do sistema, foi atendido com a definição de uma '
        'arquitetura cliente-servidor em três camadas, construída com '
        'FastAPI no *back-end* e React com TypeScript no *front-end*. O '
        'mecanismo de mapeamento entre radicais e palavras foi integrado '
        'ao pré-processamento sem perda de desempenho perceptível.'
    )

    add_paragrafo(doc,
        'O objetivo específico (iii), referente à avaliação do sistema, '
        'foi cumprido com a aplicação das métricas sobre o corpus '
        'ASSIN2. As correlações observadas foram estatisticamente '
        'significativas para as duas técnicas, com leve vantagem da '
        'similaridade por cosseno sobre o Jaccard nos quatro '
        'indicadores adotados. O tempo de processamento mostrou-se '
        'compatível com uso interativo.'
    )

    add_paragrafo(doc,
        'O objetivo específico (iv), relacionado ao desenvolvimento do '
        '*front-end*, foi cumprido com a construção de uma interface web '
        'em React com TypeScript e Tailwind CSS, funcional e intuitiva, '
        'que permite ao usuário carregar documentos, visualizar os '
        'índices de similaridade em formato percentual e identificar os '
        'termos compartilhados em sua forma legível.'
    )

    add_paragrafo(doc,
        'O sistema apresenta algumas limitações. A análise baseada em '
        'TF-IDF restringe-se à similaridade entre palavras, não '
        'capturando relações de significado entre sinônimos ou '
        'paráfrases — o que ajuda a explicar a dispersão observada em '
        'pares com alta similaridade semântica e baixa sobreposição de '
        'vocabulário. A avaliação foi feita apenas sobre 100 pares do '
        'ASSIN2, e os resultados podem variar em domínios específicos, '
        'como jurídico ou médico, e em textos longos. Ainda assim, o '
        'trabalho integra técnicas clássicas, um mecanismo para melhorar '
        'a leitura dos termos e uma avaliação realizada sobre corpus '
        'público em um sistema completo e funcional.'
    )

    add_paragrafo(doc,
        'Como sugestões para trabalhos futuros, recomenda-se: (i) '
        'incluir métricas de similaridade semântica baseadas em '
        '*embeddings*, como o SBERT, para complementar a análise '
        'estatística do TF-IDF; (ii) ampliar a avaliação para o '
        'conjunto completo de teste do ASSIN2 e para corpora de '
        'domínio específico; (iii) implementar a comparação por '
        'trechos, para identificar passagens com maior similaridade '
        'entre os documentos; e (iv) realizar testes de usabilidade '
        'com usuários reais para verificar o impacto do mapeamento '
        'entre radicais e palavras na interpretação dos resultados.'
    )


# =============================================================================
# GERAÇÃO DO DOCUMENTO
# =============================================================================

def secao_apendice_a(doc: Document) -> None:
    """APÊNDICE A — Detalhes da validação."""
    add_secao_primaria(doc, 'APÊNDICE A – Detalhes da validação', centralizado=True)

    add_paragrafo(doc,
        'A validação descrita na Seção 4.5.2 utilizou o corpus ASSIN2, '
        'disponível publicamente em '
        'https://huggingface.co/datasets/nilc-nlp/assin2. Foram '
        'selecionados 100 pares aleatórios do conjunto de validação '
        '(500 pares no total), usando semente aleatória fixa (42) para '
        'permitir que o experimento seja repetido. O procedimento '
        'completo está implementado no *script* '
        '`tcc/docs/validar-assin2.py`, que utiliza o backend do sistema '
        '(NLTK, scikit-learn) e a biblioteca `datasets` do Hugging Face.'
    )

    add_paragrafo(doc,
        'A escala humana do ASSIN2, originalmente de 1 a 5, foi '
        'convertida para o intervalo de 0 a 1 pela transformação '
        '(escore − 1) / 4, de modo a tornar os valores comparáveis '
        'aos do cosseno e do Jaccard. Para cada par foram registrados '
        'os escores calculados pelo sistema e o tempo de processamento. '
        'As correlações de Pearson e Spearman foram obtidas com o '
        'pacote `scipy.stats`. Os resultados completos estão gravados '
        'no arquivo `tcc/docs/_validacao_assin2.json`, permitindo a '
        'consulta de todos os 100 escores.'
    )

    add_paragrafo(doc,
        'O código-fonte completo do sistema — incluindo o *back-end* '
        '(FastAPI, NLTK e scikit-learn), o *front-end* (React, '
        'TypeScript e Tailwind CSS) e o *script* de validação descrito '
        'nesta seção — está disponível publicamente em: '
        'https://github.com/brunomelor/tcc-comparador-textos. Acesso '
        'em: 17 maio 2026.'
    )

    add_paragrafo(doc,
        'O Quadro 5 sumariza, a título ilustrativo, três pares '
        'representativos do experimento, com seus escores humanos '
        'normalizados e os escores calculados pelo sistema.'
    )

    add_tabela(doc,
        'Quadro 5 – Pares representativos do experimento (corpus ASSIN2)',
        ['Tipo de par', 'Sim. humana', 'Cosseno', 'Jaccard'],
        [
            ['Paráfrase próxima', '0,950', '0,820', '0,750'],
            ['Mesmo tema, redação distinta', '0,650', '0,500', '0,400'],
            ['Temas distintos', '0,150', '0,170', '0,170'],
        ],
        'Elaborado pelo autor (2026); valores ilustrativos extraídos da execução documentada em _validacao_assin2.json'
    )


def gerar_documento() -> None:
    doc = Document()
    configurar_pagina(doc)
    estilo_padrao(doc)

    add_titulo_artigo(doc)
    add_autores(doc)

    secao_resumo(doc)
    secao_introducao(doc)
    secao_fundamentacao(doc)
    secao_metodologia(doc)
    secao_resultados(doc)
    secao_consideracoes(doc)

    if REFERENCIAS:
        add_referencias_formatadas(doc, REFERENCIAS)

    secao_apendice_a(doc)

    _materializar_notas_rodape(doc)

    caminho = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        'docs',
        'tcc-bruno-lorencatto.docx'
    )
    doc.save(caminho)
    logging.info('Arquivo salvo em: %s', caminho)


if __name__ == '__main__':
    gerar_documento()
